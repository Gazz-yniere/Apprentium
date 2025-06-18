from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_DIRECTION
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph # Import pour pouvoir wrapper lxml element
from docx.oxml import OxmlElement
from calculs_generator import generate_arithmetic_problems # Modifié pour utiliser la nouvelle fonction
from pdf_generator import SECTION_ASSETS, get_resource_path_pdf # Pour les images et couleurs (get_resource_path_pdf est ok ici)
import random
import warnings
import os

# Ignorer l'avertissement spécifique de python-docx concernant la recherche de style
warnings.filterwarnings("ignore", message="style lookup by style_id is deprecated")

def get_output_path(filename, custom_output_dir=None):
    import sys, os
    base_dir_for_output = None

    if custom_output_dir and custom_output_dir.strip():
        custom_output_dir = custom_output_dir.strip()
        if os.path.isdir(custom_output_dir):
            base_dir_for_output = os.path.normpath(custom_output_dir)
        else:
            try:
                os.makedirs(custom_output_dir, exist_ok=True)
                base_dir_for_output = os.path.normpath(custom_output_dir)
            except OSError:
                print(f"Avertissement : Impossible de créer le dossier personnalisé '{custom_output_dir}'. Utilisation du dossier par défaut.")

    if base_dir_for_output is None: # Si custom_output_dir n'est pas fourni ou a échoué
        if getattr(sys, 'frozen', False): # Exécutable PyInstaller
            app_base_path = os.path.dirname(sys.executable)
        else: # Mode script
            app_base_path = os.path.dirname(os.path.dirname(__file__)) # Remonte au dossier parent de src/
        base_dir_for_output = os.path.normpath(os.path.join(app_base_path, 'output'))

    os.makedirs(base_dir_for_output, exist_ok=True) # S'assure que le dossier final existe
    return os.path.join(base_dir_for_output, filename)

def add_math_problems_to_doc(container, problems_for_day, indent_val=None):
    if not problems_for_day:
        return
    
    # Ajustement du titre en fonction du nombre de problèmes
    if len(problems_for_day) == 1:
        title_str = "Petit Problème:"
    else:
        title_str = "Petits Problèmes:"
    # Remplacer add_heading par add_paragraph et styler le run pour le titre
    p_heading = container.add_paragraph()
    run_heading = p_heading.add_run(title_str)
    run_heading.bold = True
    # run_heading.underline = True # Retiré pour correspondre aux titres d'opérations
    # run_heading.font.size = Pt(12) # Retiré pour utiliser la taille par défaut (11pt) comme les titres d'opérations
    if indent_val:
        p_heading.paragraph_format.left_indent = indent_val
    p_heading.paragraph_format.space_before = Pt(6)
    p_heading.paragraph_format.space_after = Pt(3)

    for idx, problem_data in enumerate(problems_for_day):
        problem_text = f"{idx + 1}. {problem_data['content']}"
        p_problem = container.add_paragraph()
        # Numéro en gras
        run_number = p_problem.add_run(f"{idx + 1}. ")
        run_number.bold = True
        # Texte de l'énoncé en normal
        p_problem.add_run(problem_data['content'])

        if indent_val:
            p_problem.paragraph_format.left_indent = indent_val
        p_problem.paragraph_format.space_before = Pt(0)
        p_problem.paragraph_format.space_after = Pt(1)

        p_answer = container.add_paragraph("Réponse: ______________________________")
        if indent_val:
            p_answer.paragraph_format.left_indent = indent_val
        p_answer.paragraph_format.space_before = Pt(0)
        
        if idx < len(problems_for_day) - 1:
            p_answer.paragraph_format.space_after = Pt(6) # Espace avant le prochain problème
        else:
            p_answer.paragraph_format.space_after = Pt(1) # Espace normal après le dernier problème
             
def set_table_borders_invisible(table):
    """Rend toutes les bordures d'un tableau invisibles."""
    tbl = table._tbl
    tblPr = tbl.tblPr # Get table properties element
    if tblPr is None: # Should not happen with add_table, but good practice
        tblPr = OxmlElement('w:tblPr')
        tbl.append(tblPr)

    # Vérifie si tblBorders existe déjà, sinon le crée
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
    else: # Clear existing border elements to ensure a clean slate
        for child in list(tblBorders):
            tblBorders.remove(child)

    # Définit tous les types de bordures à 'nil'
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'nil')
        # Optionnel: pour s'assurer qu'elles sont vraiment invisibles, on peut aussi mettre size, space, color à 0/auto
        border_el.set(qn('w:sz'), '0')
        border_el.set(qn('w:space'), '0')
        border_el.set(qn('w:color'), 'auto')
        tblBorders.append(border_el)

def set_table_borders_visible_colored(table, rgb_float_color, size_pt=6):
    """Applique des bordures visibles colorées à un tableau (pour l'encadrement de section)."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.append(tblPr)

    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
    else:
        for child in list(tblBorders): # Nettoyer les bordures existantes
            tblBorders.remove(child)
    
    hex_color_str = "{:02x}{:02x}{:02x}".format(
        int(rgb_float_color[0]*255), 
        int(rgb_float_color[1]*255), 
        int(rgb_float_color[2]*255)
    )

    for border_name in ['top', 'left', 'bottom', 'right']: # Bordures extérieures uniquement
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'single') 
        border_el.set(qn('w:sz'), str(size_pt))  # Taille en 1/8 de point (ex: 6 pour 0.75pt)
        border_el.set(qn('w:space'), '0')
        border_el.set(qn('w:color'), hex_color_str)
        tblBorders.append(border_el)

def add_section_header_word(parent_cell, section_key_name, section_num_val):
    """Ajoute un en-tête de section stylisé avec titre et image au document Word."""
    section_data = SECTION_ASSETS.get(section_key_name, {})
    color_tuple_float = section_data.get("color", (0,0,0))
    docx_section_color = RGBColor(
        int(color_tuple_float[0] * 255),
        int(color_tuple_float[1] * 255),
        int(color_tuple_float[2] * 255)
    )
    image_path_relative = section_data.get("image_path")

    # Le header_table est maintenant ajouté à la cellule parente
    header_table = parent_cell.add_table(rows=1, cols=2)
    # Ses propres bordures doivent être invisibles car il est dans un cadre plus grand
    set_table_borders_invisible(header_table)
    header_table.autofit = False
    header_table.allow_autofit = False

    # Largeur cible pour le tableau d'en-tête de section, environ la moitié de la largeur disponible.
    # Contenu disponible approx. 6.87 pouces. Moitié = approx 3.4 pouces.
    target_table_width = Inches(3.4)

    # Ajuster les largeurs pour l'image et le titre
    image_column_width = Inches(0.5)
    # La colonne du titre prend le reste de la largeur CIBLE du tableau
    header_table.columns[0].width = target_table_width - image_column_width
    header_table.columns[1].width = image_column_width

    # Cellule du titre
    title_cell = header_table.cell(0, 0)
    title_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p_title = title_cell.paragraphs[0]
    p_title.clear() # Nettoyer le paragraphe par défaut
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(0)

    run_num = p_title.add_run(f"{section_num_val}. ")
    run_num.bold = True
    run_num.font.size = Pt(14)
    run_num.font.color.rgb = docx_section_color

    run_title_text = p_title.add_run(section_key_name)
    run_title_text.bold = True
    run_title_text.font.size = Pt(14)
    run_title_text.font.color.rgb = docx_section_color

    # Cellule de l'image
    image_cell = header_table.cell(0, 1)
    image_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP # Aligner l'image en haut
    img_para = image_cell.paragraphs[0]
    img_para.clear()
    img_para.paragraph_format.space_before = Pt(0)
    img_para.paragraph_format.space_after = Pt(0)
    img_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if image_path_relative:
        image_path_abs = get_resource_path_pdf(image_path_relative) # Utilise la fonction de pdf_generator
        if image_path_abs and os.path.exists(image_path_abs):
            try:
                img_para.add_run().add_picture(image_path_abs, width=Inches(0.4)) # Taille de l'image réduite
            except Exception as e:
                print(f"Erreur add_picture (Word) pour {image_path_abs}: {e}")

def delete_paragraph(paragraph):
    """Supprime un paragraphe du document."""
    # Cette fonction n'est plus utilisée directement pour supprimer les paragraphes auto-générés
    # mais est conservée au cas où elle serait utilisée ailleurs.
    # La logique de suppression des paragraphes vides auto-générés est maintenant dans delete_paragraph_by_element.
    if paragraph is None or paragraph._element is None or paragraph._element.getparent() is None:
        return
    p_element = paragraph._element
    p_element.getparent().remove(p_element)

def delete_paragraph_by_element(p_element, parent_docx_object):
    """Supprime un paragraphe en utilisant son élément lxml et l'objet parent python-docx,
       uniquement s'il semble être un paragraphe vide auto-généré."""
    if p_element is None or p_element.getparent() is None: return False
    try:
        p = Paragraph(p_element, parent_docx_object)
        if not p.text.strip() and not p.runs:
            p_element.getparent().remove(p_element)
            return True 
    except Exception as e:
        print(f"Erreur lors de la vérification/suppression du paragraphe par élément : {e}")
        pass
    return False

def set_cell_margins(cell, top=None, bottom=None, left=None, right=None):
    """Définit les marges internes d'une cellule en points (Pt)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Chercher l'élément <w:tcMar> existant.
    # tcPr est un objet OxmlElement (CT_TcPr).
    tcMar_element = tcPr.find(qn('w:tcMar'))
    if tcMar_element is None:
        # S'il n'existe pas, le créer et l'ajouter à tcPr.
        tcMar_element = OxmlElement('w:tcMar')
        tcPr.append(tcMar_element)
    # tcMar_element est maintenant l'élément lxml <w:tcMar>.

    # Définir les marges individuelles
    margins_to_set = {
        "top": top, "bottom": bottom, "left": left, "right": right
    }

    for margin_name, value_pt in margins_to_set.items():
        if value_pt is not None:
            # Chercher l'élément de marge spécifique (ex: <w:top>) dans <w:tcMar>
            margin_sub_element = tcMar_element.find(qn(f'w:{margin_name}'))
            if margin_sub_element is None:
                margin_sub_element = OxmlElement(f'w:{margin_name}')
                tcMar_element.append(margin_sub_element)
            margin_sub_element.set(qn("w:w"), str(int(value_pt.pt * 20))) # Convertir points en DXA (twips)
            margin_sub_element.set(qn("w:type"), "dxa")


def generate_workbook_docx(days, operations, counts, max_digits, conjugations, params_list, grammar_exercises,
                           orthographe_exercises=None, enumerate_exercises=None, sort_exercises=None,
                           geo_exercises=None, english_exercises=None, encadrement_exercises_list=None, # Modifié
                           story_math_problems_by_day=None, # Ajout du paramètre
                           compare_numbers_exercises_list=None, # Nouveau
                           logical_sequences_exercises_list=None, # Nouveau
                           header_text=None, show_name=False, show_note=False, filename="workbook.docx", output_dir_override=None):
    if geo_exercises is None:
        geo_exercises = []
    if english_exercises is None:
        english_exercises = []
    if orthographe_exercises is None: orthographe_exercises = []
    if enumerate_exercises is None: enumerate_exercises = []
    if sort_exercises is None: sort_exercises = []
    if encadrement_exercises_list is None: encadrement_exercises_list = [] # Modifié
    if story_math_problems_by_day is None: # Initialisation
        story_math_problems_by_day = []

    doc = Document()
    # Style général
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0, 0, 0)
    # Réduction des marges haut et bas
    doc_section = doc.sections[0]
    doc_section.top_margin = Inches(0.4)
    doc_section.bottom_margin = Inches(0.4)
    doc_section.left_margin = Inches(0.7) # Un peu plus de marge à gauche pour l'ensemble
    doc_section.right_margin = Inches(0.7)

    CONTENT_INDENT = Cm(0.8) # Décalage pour le contenu des sections

    # Applique l'espacement après à tous les paragraphes créés
    def add_paragraph(parent_container, text="", style=None, indent=False):
        para = parent_container.add_paragraph(text, style=style) if style else parent_container.add_paragraph(text)
        para.paragraph_format.space_before = Pt(0) # Assurer aucun espace avant
        para.paragraph_format.space_after = Cm(0.05)
        if indent:
            para.paragraph_format.left_indent = CONTENT_INDENT
        return para

    # En-tête général du document (Nom/Titre/Note)
    # Ce tableau n'est PAS dans un cadre de section coloré.
    # Sa largeur est gérée par python-docx pour s'adapter aux marges.
    for day in range(1, days + 1):
        # En-tête
        if header_text or show_name or show_note:
            table = doc.add_table(rows=1, cols=3)
            table.autofit = False
            row = table.rows[0]
            if show_name:
                cell = row.cells[0]
                p = cell.paragraphs[0]
                run = p.add_run("Nom : ____________________")
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                row.cells[0].text = ""
            cell = row.cells[1]
            p = cell.paragraphs[0]
            run = p.add_run(header_text or "")
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if show_note:
                cell = row.cells[2]
                p = cell.paragraphs[0]
                run = p.add_run("Note : _______")
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                row.cells[2].text = ""
            add_paragraph(doc, "") # Espace après l'en-tête général, ajouté au document principal
        
        # Préparer les données des "Petits Problèmes" pour le jour
        current_day_story_problems = None
        if len(story_math_problems_by_day) >= day:
            current_day_story_problems = story_math_problems_by_day[day-1]

        section_num = 1 # Compteur pour la numérotation des sections

        # Section Calculs
        if any(counts) or \
           (enumerate_exercises and len(enumerate_exercises) >= day and enumerate_exercises[day-1]) or \
           bool(current_day_story_problems): # Condition pour inclure les "Petits Problèmes"
            section_key = "Calculs"
            section_frame_table = doc.add_table(rows=1, cols=1)
            section_color_data = SECTION_ASSETS.get(section_key, {}).get("color", (0.1,0.1,0.1))
            set_table_borders_visible_colored(section_frame_table, section_color_data, size_pt=8) # Cadre visible
            
            section_cell = section_frame_table.cell(0,0)
            set_cell_margins(section_cell, top=Pt(0), bottom=Pt(10), left=Cm(0.15), right=Cm(0.15))
            if section_cell.paragraphs: # Nettoyer le paragraphe par défaut
                p_default = section_cell.paragraphs[0]
                if not p_default.text and not p_default.runs: delete_paragraph(p_default)
            
            add_section_header_word(section_cell, section_key, section_num)
            section_num += 1
            
            # Tenter de supprimer le paragraphe vide par défaut ajouté après le tableau d'en-tête
            if section_cell.tables:
                header_table_element = section_cell.tables[0]._tbl 
                tc_element = section_cell._tc
                children = list(tc_element) # Accéder directement aux enfants de l'élément CT_Tc
                try:
                    table_index = children.index(header_table_element) # header_table_element est déjà l'élément XML
                    if table_index + 1 < len(children) and children[table_index + 1].tag == qn('w:p'):
                         delete_paragraph_by_element(children[table_index + 1], section_cell)
                except ValueError:
                    pass 

            # Enumérer un nombre (si présent pour le jour)
            if enumerate_exercises and len(enumerate_exercises) >= day and enumerate_exercises[day-1]:
                para_enum = add_paragraph(section_cell, indent=True)
                run_enum = para_enum.add_run("Écris chaque nombre en toutes lettres :")
                run_enum.bold = True
                for n_enum in enumerate_exercises[day-1]:
                    add_paragraph(section_cell, f"{n_enum} = _____________________________________________", style='ListContinue', indent=True)

            # Opérations classiques
            if any(counts):
                for i, operation in enumerate(operations):
                    # Vérifier si cette opération spécifique a des exercices
                    if params_list[i].get('count', 0) > 0:
                        p_op = add_paragraph(section_cell, indent=True) 
                        run_op = p_op.add_run(f"{operation.capitalize()} :")
                        run_op.bold = True
                        problems = generate_arithmetic_problems(operation, params_list[i]) # Utilise la fonction importée
                        for problem in problems:
                            calc_str = problem.strip().replace(' =', '')
                            add_paragraph(section_cell, f"{calc_str} = ________________________________", style='ListContinue', indent=True)
            
            
            # Ajout des "Petits Problèmes" à la fin de la section Calculs
            if current_day_story_problems:
                add_math_problems_to_doc(section_cell, current_day_story_problems, indent_val=CONTENT_INDENT)
                # Pas besoin d'un add_paragraph(section_cell, "") ici, car add_math_problems_to_doc gère l'espacement interne
                # et la cellule a une marge inférieure.
            add_paragraph(doc, "") # Espace entre les cadres de section

        if conjugations and len(conjugations) >= day and conjugations[day-1]:
            section_key = "Conjugaison"
            section_frame_table = doc.add_table(rows=1, cols=1)
            section_color_data = SECTION_ASSETS.get(section_key, {}).get("color", (0.1,0.1,0.1))
            set_table_borders_visible_colored(section_frame_table, section_color_data, size_pt=8)
            section_cell = section_frame_table.cell(0,0)
            set_cell_margins(section_cell, top=Pt(0), bottom=Pt(10), left=Cm(0.15), right=Cm(0.15))
            if section_cell.paragraphs: 
                p_default = section_cell.paragraphs[0]
                if not p_default.text and not p_default.runs: delete_paragraph(p_default)

            add_section_header_word(section_cell, section_key, section_num)
            section_num += 1

            # Tenter de supprimer le paragraphe vide par défaut ajouté après le tableau d'en-tête
            if section_cell.tables:
                header_table_element = section_cell.tables[0]._tbl
                tc_element = section_cell._tc
                children = list(tc_element)
                try:
                    table_index = children.index(header_table_element)
                    if table_index + 1 < len(children) and children[table_index + 1].tag == qn('w:p'):
                         delete_paragraph_by_element(children[table_index + 1], section_cell)
                except ValueError:
                    pass

            from conjugation_generator import PRONOUNS, VERBS
            for conjugation_item in conjugations[day-1]:
                verb = conjugation_item["verb"]
                tense = conjugation_item["tense"]
                groupe = None
                for g in (1, 2, 3):
                    if verb in VERBS.get(g, []): 
                        groupe = f"{g}er groupe"
                        break
                if not groupe and "usuels" in VERBS and verb in VERBS["usuels"]:
                    groupe = "usuel"
                if not groupe: groupe = "inconnu" # Fallback

                p_conj_details = add_paragraph(section_cell, indent=True)
                run = p_conj_details.add_run(f"Verbe : {verb}  |  Groupe : {groupe}  |  Temps : {tense}")
                run.bold = True
                for pronoun in PRONOUNS:
                    add_paragraph(section_cell, f"{pronoun} ____________________", indent=True)
            add_paragraph(doc, "") # Espace entre les cadres de section

        if grammar_exercises and len(grammar_exercises) >= day and grammar_exercises[day-1]:
            section_key = "Grammaire"
            section_frame_table = doc.add_table(rows=1, cols=1)
            section_color_data = SECTION_ASSETS.get(section_key, {}).get("color", (0.1,0.1,0.1))
            set_table_borders_visible_colored(section_frame_table, section_color_data, size_pt=8)
            section_cell = section_frame_table.cell(0,0)
            set_cell_margins(section_cell, top=Pt(0), bottom=Pt(10), left=Cm(0.15), right=Cm(0.15))
            if section_cell.paragraphs:
                p_default = section_cell.paragraphs[0]
                if not p_default.text and not p_default.runs: delete_paragraph(p_default)
            
            add_section_header_word(section_cell, section_key, section_num)
            section_num += 1

            # Tenter de supprimer le paragraphe vide par défaut ajouté après le tableau d'en-tête
            if section_cell.tables:
                header_table_element = section_cell.tables[0]._tbl
                tc_element = section_cell._tc
                children = list(tc_element)
                try:
                    table_index = children.index(header_table_element)
                    if table_index + 1 < len(children) and children[table_index + 1].tag == qn('w:p'):
                         delete_paragraph_by_element(children[table_index + 1], section_cell)
                except ValueError:
                    pass

            for ex in grammar_exercises[day-1]:
                phrase_content = ex['phrase']
                transformation_content = ex['transformation']

                phrase_line_para = add_paragraph(section_cell, indent=True)
                phrase_label_run = phrase_line_para.add_run("Phrase : ")
                phrase_label_run.bold = True
                phrase_line_para.add_run(phrase_content)

                transfo_line_para = add_paragraph(section_cell, indent=True)
                transfo_label_run = transfo_line_para.add_run("Transformation demandée : ")
                transfo_label_run.bold = True
                transfo_line_para.add_run(transformation_content)

                add_paragraph(section_cell, "Réponse : __________________________________________________________", indent=True)
            add_paragraph(doc, "") # Espace entre les cadres de section

        has_mesures_content_for_day = False
        if geo_exercises and len(geo_exercises) >= day and geo_exercises[day-1]:
            has_mesures_content_for_day = True
        if sort_exercises and len(sort_exercises) >= day and sort_exercises[day-1]:
            has_mesures_content_for_day = True

        current_day_encadrement_lines = encadrement_exercises_list[day-1] if len(encadrement_exercises_list) >= day else None
        if current_day_encadrement_lines: has_mesures_content_for_day = True
        current_day_compare_numbers = compare_numbers_exercises_list[day-1] if compare_numbers_exercises_list and len(compare_numbers_exercises_list) >= day else None
        if current_day_compare_numbers: has_mesures_content_for_day = True
        current_day_logical_sequences = logical_sequences_exercises_list[day-1] if logical_sequences_exercises_list and len(logical_sequences_exercises_list) >= day else None
        if current_day_logical_sequences: has_mesures_content_for_day = True
        
        if has_mesures_content_for_day:
            section_key = "Mesures"
            section_frame_table = doc.add_table(rows=1, cols=1)
            section_color_data = SECTION_ASSETS.get(section_key, {}).get("color", (0.1,0.1,0.1))
            set_table_borders_visible_colored(section_frame_table, section_color_data, size_pt=8)
            section_cell = section_frame_table.cell(0,0)
            set_cell_margins(section_cell, top=Pt(0), bottom=Pt(10), left=Cm(0.15), right=Cm(0.15))
            if section_cell.paragraphs:
                p_default = section_cell.paragraphs[0]
                if not p_default.text and not p_default.runs: delete_paragraph(p_default)

            add_section_header_word(section_cell, section_key, section_num)
            section_num += 1

            # Tenter de supprimer le paragraphe vide par défaut ajouté après le tableau d'en-tête
            if section_cell.tables:
                header_table_element = section_cell.tables[0]._tbl
                tc_element = section_cell._tc
                children = list(tc_element)
                try:
                    table_index = children.index(header_table_element)
                    if table_index + 1 < len(children) and children[table_index + 1].tag == qn('w:p'):
                         delete_paragraph_by_element(children[table_index + 1], section_cell)
                except ValueError:
                    pass

            if geo_exercises and len(geo_exercises) >= day and geo_exercises[day-1]:
                current_day_geo_ex = geo_exercises[day-1]
                if current_day_geo_ex:
                    para_conv_title = add_paragraph(section_cell, indent=True)
                    # Ajustement du titre en fonction du nombre d'exercices de conversion
                    if len(current_day_geo_ex) == 1:
                        run_conv_title = para_conv_title.add_run("Conversion :")
                    else:
                        run_conv_title = para_conv_title.add_run("Conversions :")
                    run_conv_title.bold = True
                    for ex_conv in current_day_geo_ex:
                        add_paragraph(section_cell, ex_conv, indent=True)
            if sort_exercises and len(sort_exercises) >= day and sort_exercises[day-1]:
                current_day_sort_ex = sort_exercises[day-1]
                if current_day_sort_ex:
                    para_sort_title = add_paragraph(section_cell, indent=True)
                    ordre = "croissant" if current_day_sort_ex[0]['type'] == 'croissant' else "décroissant"
                    run_sort_title = para_sort_title.add_run(f"Range les nombres suivants dans l'ordre {ordre} :")
                    run_sort_title.bold = True
                    for ex_sort in current_day_sort_ex:
                        numbers_str = ", ".join(str(n) for n in ex_sort['numbers'])
                        add_paragraph(section_cell, f"{numbers_str} = _________________________________", indent=True)
            if current_day_encadrement_lines:
                para_enc_title = add_paragraph(section_cell, indent=True)
                # Ajustement du titre en fonction du nombre d'exercices d'encadrement
                if len(current_day_encadrement_lines) == 1:
                    run_enc_title = para_enc_title.add_run("Encadre le nombre :")
                else:
                    run_enc_title = para_enc_title.add_run("Encadre les nombres :")
                run_enc_title.bold = True
                for ex_enc in current_day_encadrement_lines:
                    n_enc, t_enc = ex_enc['number'], ex_enc['type']
                    label = f"à l'{t_enc}" if t_enc == "unité" else f"à la {t_enc}" if t_enc in ["dizaine", "centaine"] else f"au {t_enc}"
                    add_paragraph(section_cell, f"{n_enc} {label} : ______  {n_enc}  ______", indent=True)
            add_paragraph(doc, "") # Espace entre les cadres de section
            
            # Intégration de "Comparer des nombres" dans la section Mesures (Word)
            if current_day_compare_numbers:
                para_compare_title = add_paragraph(section_cell, indent=True)
                run_compare_title = para_compare_title.add_run("Comparer les nombres (<, >, =) :")
                run_compare_title.bold = True
                for ex_compare in current_day_compare_numbers:
                    add_paragraph(section_cell, f"{ex_compare['num1']} ______ {ex_compare['num2']}", indent=True)

            # Intégration de "Suites Logiques" dans la section Mesures (Word)
            if current_day_logical_sequences:
                para_seq_title = add_paragraph(section_cell, indent=True)
                run_seq_title = para_seq_title.add_run("Complète les suites logiques :")
                run_seq_title.bold = True
                for ex_seq in current_day_logical_sequences:
                    sequence_str = " ".join(map(str, ex_seq['sequence_displayed']))
                    # type_desc = ""
                    # if ex_seq['type'] == 'arithmetic_plus':
                    #     type_desc = f" (+{ex_seq['step']})"
                    # elif ex_seq['type'] == 'arithmetic_minus':
                    #     type_desc = f" (-{ex_seq['step']})"
                    add_paragraph(section_cell, f"{sequence_str}", indent=True)
            # L'espace après la section Mesures est géré par add_paragraph(doc, "") plus bas si d'autres sections suivent

        if english_exercises and len(english_exercises) >= day and english_exercises[day-1]: 
            section_key = "Anglais"
            section_frame_table = doc.add_table(rows=1, cols=1)
            section_color_data = SECTION_ASSETS.get(section_key, {}).get("color", (0.1,0.1,0.1))
            set_table_borders_visible_colored(section_frame_table, section_color_data, size_pt=8)
            section_cell = section_frame_table.cell(0,0)
            set_cell_margins(section_cell, top=Pt(0), bottom=Pt(10), left=Cm(0.15), right=Cm(0.15))
            if section_cell.paragraphs:
                p_default = section_cell.paragraphs[0]
                if not p_default.text and not p_default.runs: delete_paragraph(p_default)
            
            add_section_header_word(section_cell, section_key, section_num)
            section_num += 1

            # Tenter de supprimer le paragraphe vide par défaut ajouté après le tableau d'en-tête
            if section_cell.tables:
                header_table_element = section_cell.tables[0]._tbl
                tc_element = section_cell._tc
                children = list(tc_element)
                try:
                    table_index = children.index(header_table_element)
                    if table_index + 1 < len(children) and children[table_index + 1].tag == qn('w:p'):
                         delete_paragraph_by_element(children[table_index + 1], section_cell)
                except ValueError:
                    pass

            completer_subtitle_shown = False
            relier_subtitle_shown = False # Non utilisé, mais conservé pour la logique

            for ex in english_exercises[day-1]:
                if ex['type'] in ('simple', 'complexe'):
                    if not completer_subtitle_shown:
                        para_complete_title = add_paragraph(section_cell, indent=True)
                        run_complete_title = para_complete_title.add_run("Compléter :")
                        run_complete_title.bold = True
                        completer_subtitle_shown = True
                    add_paragraph(section_cell, ex['content'], indent=True)
                elif ex['type'] == 'relier':
                    para_relier_title = add_paragraph(section_cell, indent=True)
                    run_relier_title = para_relier_title.add_run("Jeu de mots à relier :")
                    run_relier_title.bold = True
                    
                    mots_a_relier_pour_jeu = ex['content']
                    liste_anglais = [item['english'] for item in mots_a_relier_pour_jeu]
                    liste_francais = [item['french'] for item in mots_a_relier_pour_jeu]
                    random.shuffle(liste_francais) 
                    max_items = len(liste_anglais)
                    
                    if max_items > 0:
                        # Créer un seul tableau pour tout le jeu à relier
                        table = section_cell.add_table(rows=max_items, cols=6) # Pré-allouer les lignes
                        table.autofit = False 
                        table.allow_autofit = False 
                        tblPr = table._tbl.tblPr

                        # Explicitly set table layout to fixed to enforce column widths
                        tblLayout_obj = tblPr.get_or_add_tblLayout() # Gets or adds <w:tblLayout>
                        tblLayout_obj.type = 'fixed' # Sets <w:tblLayout w:type="fixed"/>

                        col_indent_width = CONTENT_INDENT # Largeur de la nouvelle colonne d'indentation
                        col_word_width = Inches(2) # Largeur doublée pour les colonnes de mots
                        col_bullet_width = Inches(0.1) 
                        col_space_width = Inches(0.2)  
                        
                        target_table_width_inches = col_indent_width.inches + (col_word_width.inches * 2) + (col_bullet_width.inches * 2) + col_space_width.inches
                        target_table_width_dxa = int(target_table_width_inches * 1440) 
                        tblW = OxmlElement('w:tblW')
                        tblW.set(qn('w:w'), str(target_table_width_dxa))
                        tblW.set(qn('w:type'), 'dxa') 
                        tblPr.append(tblW)
                        
                        # Définir les largeurs de colonne une seule fois pour le tableau
                        table.columns[0].width = col_indent_width   # Colonne d'indentation
                        table.columns[1].width = col_word_width     # Mot anglais
                        table.columns[2].width = col_bullet_width   # Puce 1
                        table.columns[3].width = col_space_width    # Espace central
                        table.columns[4].width = col_bullet_width   # Puce 2
                        table.columns[5].width = col_word_width     # Mot français

                        for i in range(max_items):
                            mot_anglais = liste_anglais[i]
                            mot_francais = liste_francais[i] 
                            
                            row_cells = table.rows[i].cells # Obtenir les cellules de la ligne courante
                            
                            contents = ["", mot_anglais, "\u2022", "", "\u2022", mot_francais] 
                            alignments = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, 
                                          WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT] 
                            for idx, content_text in enumerate(contents):
                                cell = row_cells[idx]
                                cell.text = content_text 
                                paragraph = cell.paragraphs[0]
                                paragraph.paragraph_format.space_before = Pt(0)
                                paragraph.paragraph_format.space_after = Pt(1) 
                                paragraph.alignment = alignments[idx]
                        
                        set_table_borders_invisible(table) # Appliquer une fois à la fin
                        # Tenter de supprimer le paragraphe vide par défaut ajouté après ce petit tableau
                        tc_element = section_cell._tc
                        current_table_element = table._tbl 
                        children = list(tc_element)
                        try:
                            table_index = children.index(current_table_element)
                            if table_index + 1 < len(children) and children[table_index + 1].tag == qn('w:p'):
                                 delete_paragraph_by_element(children[table_index + 1], section_cell)
                        except ValueError:
                            pass 
            add_paragraph(doc, "") # Espace entre les cadres de section

        if orthographe_exercises and len(orthographe_exercises) >= day and orthographe_exercises[day-1]:
            section_key = "Orthographe"
            section_frame_table = doc.add_table(rows=1, cols=1)
            section_color_data = SECTION_ASSETS.get(section_key, {}).get("color", (0.1,0.1,0.1))
            set_table_borders_visible_colored(section_frame_table, section_color_data, size_pt=8)
            section_cell = section_frame_table.cell(0,0)
            set_cell_margins(section_cell, top=Pt(0), bottom=Pt(10), left=Cm(0.15), right=Cm(0.15))
            if section_cell.paragraphs:
                p_default = section_cell.paragraphs[0]
                if not p_default.text and not p_default.runs: delete_paragraph(p_default)

            add_section_header_word(section_cell, section_key, section_num)
            section_num += 1

            # Tenter de supprimer le paragraphe vide par défaut ajouté après le tableau d'en-tête
            if section_cell.tables:
                header_table_element = section_cell.tables[0]._tbl
                tc_element = section_cell._tc
                children = list(tc_element)
                try:
                    table_index = children.index(header_table_element)
                    if table_index + 1 < len(children) and children[table_index + 1].tag == qn('w:p'):
                         delete_paragraph_by_element(children[table_index + 1], section_cell)
                except ValueError:
                    pass

            for ex_ortho in orthographe_exercises[day-1]:
                if ex_ortho['type'] == 'homophone':
                    para_homo = add_paragraph(section_cell, indent=True)
                    run_homo_type = para_homo.add_run(f"{ex_ortho['homophone']} : ")
                    run_homo_type.bold = True
                    para_homo.add_run(ex_ortho['content'])
            add_paragraph(doc, "") # Espace entre les cadres de section

        if day < days: # N'ajoute pas de saut de page après le dernier jour
            doc.add_page_break()
    out_path = get_output_path(filename, output_dir_override)
    doc.save(out_path)
    print(f"Word généré : {out_path}")
    return out_path
